"""
Parse a DOCX mock shell using python-docx to extract table text,
then send to the selected LLM for JSON structuring.
"""
import json
import io
import docx
from llm_client import call_llm
from ._shell_prompt import SHELL_PARSE_SYSTEM


def _extract_text(file_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_bytes))
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    for idx, table in enumerate(doc.tables):
        lines.append(f"\n--- TABLE {idx + 1} ---")
        for row in table.rows:
            # Preserve cell-level indentation by prefixing leading spaces from
            # the first cell's paragraph (python-docx strips visual indent),
            # so the LLM can still see the SOC > PT hierarchy in the dump.
            first = row.cells[0]
            indent = ""
            if first.paragraphs:
                p = first.paragraphs[0]
                pf = p.paragraph_format
                if pf and pf.left_indent:
                    indent = "  " * max(1, int(pf.left_indent.pt // 12))
            cells = [c.text.strip() for c in row.cells]
            if cells:
                cells[0] = indent + cells[0]
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def _strip_fences(raw: str) -> str:
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return raw.strip()


def parse_docx(
    file_bytes: bytes,
    api_key: str | None = None,
    provider: str = "OpenAI",
    model: str = "gpt-4o-mini",
) -> dict:
    extracted = _extract_text(file_bytes)
    raw = call_llm(
        system=SHELL_PARSE_SYSTEM,
        user=(
            "Parse this clinical mock shell table (extracted from Word/DOCX) "
            "into the JSON schema described in your instructions. Pay especially "
            "close attention to row indentation — leading whitespace in the row "
            "labels indicates SOC > PT (or parameter > category) nesting; every "
            "indented row must record its parent_label and an indent_level >= 1. "
            "Return ONLY the JSON object.\n\n"
            + extracted
        ),
        provider=provider,
        model=model,
        api_key=api_key or "",
        max_tokens=4000,
    )
    return json.loads(_strip_fences(raw))
