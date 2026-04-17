"""
AdaM Specifications Parser
Reads Excel / PDF / DOCX AdaM spec files and returns a structured JSON dict
that the LLM can use alongside the table shell JSON to generate accurate R code.
"""

import io
import json
from llm_client import call_llm

_SYSTEM = """
You are a clinical data standards expert. You will be given raw text extracted from an
AdaM dataset specification document. Your job is to return a structured JSON object
summarising the key programming metadata.

Return ONLY valid JSON — no markdown fences, no explanation.

Output schema:
{
  "dataset": "<dataset name e.g. ADRS>",
  "description": "<one-line description>",
  "population_flags": [{"variable": "FASFL", "condition": "FASFL='Y'", "label": "Full Analysis Set"}],
  "key_variables": [
    {"variable": "PARAMCD", "label": "Parameter Code", "type": "Char", "codelist": ["BOR", "ORR"], "notes": "..."},
    {"variable": "AVALC",   "label": "Analysis Value", "type": "Char", "codelist": ["CR","PR","SD","PD","NE"], "notes": "..."}
  ],
  "treatment_variable": "TRTP",
  "analysis_conditions": [
    {"output": "Tumor Response Table", "paramcd_filter": "BOR", "anl_flag": "ANL01FL='Y'", "primary_var": "AVALC", "derived_condition": "n(%) by AVALC", "r_skill": "group_count"}
  ],
  "codelists": [
    {"codelist": "AVALC", "code": "CR", "decode": "Complete Response"}
  ]
}

If any section is not present in the text, return an empty list [] for that field.
"""


def _strip_fences(raw: str) -> str:
    if raw.startswith("```"):
        parts = raw.split("```")
        raw = parts[1]
        if raw.startswith("json\n") or raw.startswith("JSON\n"):
            raw = raw[5:]
    return raw.strip()


def _call_llm(raw_text: str, provider: str, model: str, api_key: str) -> dict:
    text = raw_text[:40000]  # ~10k tokens — safe for all supported models (128k+ context)
    raw = call_llm(
        system=_SYSTEM,
        user=f"Extract AdaM specs JSON from the following text:\n\n{text}",
        provider=provider,
        model=model,
        api_key=api_key,
        max_tokens=4000,
        temperature=0.7,
    )
    return json.loads(_strip_fences(raw))


# ── Excel ─────────────────────────────────────────────────────────────────────
_PRIORITY_KEYWORDS = {"variable", "metadata", "value level", "codelist", "param",
                      "analysis", "derivation", "where clause", "dataset"}

def _extract_excel_text(file_bytes: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)

    # Extract all sheets, but put priority sheets (variable metadata, codelists) first
    priority_sheets = []
    other_sheets = []
    for name in wb.sheetnames:
        name_lower = name.lower()
        if any(kw in name_lower for kw in _PRIORITY_KEYWORDS):
            priority_sheets.append(name)
        else:
            other_sheets.append(name)

    lines = []
    for name in priority_sheets + other_sheets:
        ws = wb[name]
        lines.append(f"\n=== Sheet: {name} ===")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            line  = " | ".join(cells).strip(" |")
            if line:
                lines.append(line)
    return "\n".join(lines)


# ── PDF ───────────────────────────────────────────────────────────────────────
def _extract_pdf_text(file_bytes: bytes) -> str:
    import pdfplumber
    lines = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                lines.append(text)
            for tbl in page.extract_tables():
                for row in tbl:
                    lines.append(" | ".join(c or "" for c in row))
    return "\n".join(lines)


# ── DOCX ──────────────────────────────────────────────────────────────────────
def _extract_docx_text(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            lines.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(lines)


# ── Public entry points ───────────────────────────────────────────────────────
def parse_adam_excel(
    file_bytes: bytes,
    api_key: str | None = None,
    provider: str = "OpenAI",
    model: str = "gpt-4o-mini",
) -> dict:
    return _call_llm(_extract_excel_text(file_bytes), provider, model, api_key or "")


def parse_adam_pdf(
    file_bytes: bytes,
    api_key: str | None = None,
    provider: str = "OpenAI",
    model: str = "gpt-4o-mini",
) -> dict:
    return _call_llm(_extract_pdf_text(file_bytes), provider, model, api_key or "")


def parse_adam_docx(
    file_bytes: bytes,
    api_key: str | None = None,
    provider: str = "OpenAI",
    model: str = "gpt-4o-mini",
) -> dict:
    return _call_llm(_extract_docx_text(file_bytes), provider, model, api_key or "")
